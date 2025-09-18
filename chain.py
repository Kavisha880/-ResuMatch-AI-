import os
from typing import List, Dict

import streamlit as st
from dotenv import load_dotenv

from langchain_groq import ChatGroq
from langchain_core.prompts import PromptTemplate
from langchain_core.output_parsers import JsonOutputParser
from langchain_core.exceptions import OutputParserException

from docx import Document
from docx.shared import Pt, RGBColor

load_dotenv()


class Chain:
    def __init__(self):
        """
        Initialize Groq LLM safely for Streamlit Cloud.
        Reads from env first, then Streamlit Secrets.
        """
        key = os.getenv("GROQ_API_KEY") or st.secrets.get("GROQ_API_KEY")
        if not key:
            # Surface a friendly error in the UI instead of a silent crash
            st.error("GROQ_API_KEY not found. Add it in Streamlit → Settings → Secrets.")
            raise RuntimeError("Missing GROQ_API_KEY")

        # Use `model=` (more version-proof than `model_name=`)
        self.llm = ChatGroq(
            temperature=0,
            groq_api_key=key,
            model="llama-3.3-70b-versatile",
        )

    # -------------------------------
    # JD → structured JSON extractor
    # -------------------------------
    def extract_jobs(self, jd_text: str):
        """
        Extract structured job info (role, experience, skills, description) from raw JD text.
        Returns a list of dicts.
        """
        prompt_extract = PromptTemplate.from_template(
            """
            ### JOB DESCRIPTION TEXT:
            {page_data}

            ### INSTRUCTION:
            Extract the job posting and return as JSON with keys:
            - role
            - experience
            - skills
            - description

            Only return valid JSON, no extra text.
            """
        )
        chain_extract = prompt_extract | self.llm
        res = chain_extract.invoke(input={"page_data": jd_text})

        try:
            json_parser = JsonOutputParser()
            parsed = json_parser.parse(res.content)
        except OutputParserException:
            raise OutputParserException("Unable to parse JD into JSON.")

        return parsed if isinstance(parsed, list) else [parsed]

    # -------------------------------
    # Recruiter email generator
    # -------------------------------
    def write_mail(self, job, links_flat: List[Dict], user_name, user_background, user_email):
        """
        Generate a personalized cold email for the given job + portfolio links + user info.
        """
        # Turn links list into a readable inline string
        link_list = ", ".join([p.get("link", "") for p in links_flat if p.get("link")])

        prompt_email = PromptTemplate.from_template(
            """
            ### JOB DESCRIPTION:
            {job_description}

            ### USER INFO:
            Name: {user_name}
            Background: {user_background}
            Contact: {user_email}

            ### INSTRUCTION:
            Write a professional cold email to the recruiter about this role.
            - Introduce the user briefly with their background.
            - Highlight their most relevant skills based on the job description.
            - Mention these portfolio links wherever they fit: {link_list}
            - Keep tone formal but approachable.
            - Do not add preambles like "Here is the email".
            - Return only the email body.
            """
        )
        chain_email = prompt_email | self.llm
        res = chain_email.invoke({
            "job_description": str(job),
            "link_list": link_list,
            "user_name": user_name,
            "user_background": user_background,
            "user_email": user_email
        })
        return res.content

    # -------------------------------
    # Resume (.docx) generator
    # -------------------------------
    def _add_divider(self, doc):
        """Add a thin blue divider line between sections (minimal spacing)."""
        p = doc.add_paragraph("──────────────────────────────────────────────")
        run = p.runs[0]
        run.font.color.rgb = RGBColor(0, 112, 192)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)

    def write_resume(self, job, user_name, user_background, user_email, projects: List[Dict]):
        """
        Generate a one-page Resume in Word (.docx) format with:
        Summary → Education → Technical Skills → Soft Skills → Projects → Achievements

        projects: flat list like [{"name": "...", "link": "https://..."}, ...]
        """
        doc = Document()

        # --- Global style ---
        style = doc.styles['Normal']
        font = style.font
        font.size = Pt(10)

        paragraph_format = style.paragraph_format
        paragraph_format.space_after = Pt(2)
        paragraph_format.space_before = Pt(0)
        paragraph_format.line_spacing = Pt(12)

        # --- Heading styles (tight) ---
        for h in ['Heading 1', 'Heading 2', 'Heading 3']:
            h_style = doc.styles[h].paragraph_format
            h_style.space_before = Pt(2)
            h_style.space_after = Pt(2)

        # --- Bullet style (tight) ---
        bullet_style = doc.styles['List Bullet'].paragraph_format
        bullet_style.space_before = Pt(0)
        bullet_style.space_after = Pt(0)

        # --- Header ---
        doc.add_heading(user_name or "Candidate", 0)
        if user_background:
            doc.add_paragraph(user_background)
        if user_email:
            doc.add_paragraph(f"Email: {user_email}")
        self._add_divider(doc)

        # --- Summary ---
        doc.add_heading("Summary", level=1)
        summary_text = (
            "Final year B.Tech student specializing in Artificial Intelligence and Machine Learning. "
            "Hands-on experience in Generative AI, NLP, and Computer Vision projects. "
            "Passionate about applying AI to build scalable and impactful solutions."
        )
        doc.add_paragraph(summary_text)
        self._add_divider(doc)

        # --- Education ---
        doc.add_heading("Education", level=1)
        doc.add_paragraph("⚠️ Fill this section manually — Degree, College, Year, CGPA")
        self._add_divider(doc)

        # --- Technical Skills ---
        doc.add_heading("Technical Skills", level=1)
        doc.add_paragraph(
            "Python | SQL | C++ | TensorFlow | PyTorch | LangChain | Streamlit | ChromaDB | PowerBI | Git"
        )
        self._add_divider(doc)

        # --- Soft Skills ---
        doc.add_heading("Soft Skills", level=1)
        doc.add_paragraph("Communication | Adaptability | Problem-Solving | Teamwork | Leadership")
        self._add_divider(doc)

        # --- Projects ---
        doc.add_heading("Projects", level=1)
        if not projects:
            doc.add_paragraph("No matching projects found. Add portfolio links or GitHub username.")
        else:
            for project in projects[:3]:  # limit to 2–3 projects
                project_name = project.get("name") or "Untitled Project"
                project_link = project.get("link") or ""

                doc.add_paragraph(f"{project_name} ({project_link})", style="Heading 3")

                # Generate exactly 2 bullets with LLM (short & ATS-friendly)
                desc_prompt = f"""
                Write exactly 2 concise bullet points (no preamble, no numbering)
                describing the project '{project_name}'. If helpful, use context: {project_link}.
                Keep each bullet short, impactful, and ATS-friendly.
                """
                desc_chain = PromptTemplate.from_template(desc_prompt) | self.llm
                desc_res = desc_chain.invoke({})
                bullets = [b.strip("•- ") for b in desc_res.content.split("\n") if b.strip()]

                if len(bullets) < 2:
                    bullets += [
                        "Built and deployed clean, modular components end-to-end.",
                        "Improved performance and reliability through measurement and iteration."
                    ][:2 - len(bullets)]

                for b in bullets[:2]:
                    doc.add_paragraph(b, style="List Bullet")

        self._add_divider(doc)

        # --- Achievements ---
        doc.add_heading("Achievements", level=1)
        doc.add_paragraph(
            "• Solved 300+ problems on LeetCode, improving algorithmic thinking.\n"
            "• 5★ Python programmer on HackerRank.\n"
            "• Principal’s Excellence Award for AI/ML project presentation (2024).\n"
            "• Contributed to multiple open-source projects on GitHub."
        )

        # Save File
        file_name = f"{(user_name or 'Candidate').replace(' ', '_')}_Resume.docx"
        doc.save(file_name)
        return file_name
